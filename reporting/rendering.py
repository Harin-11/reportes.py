from __future__ import annotations

import argparse
import copy
import unicodedata
from pathlib import Path

from docx.shared import Pt

from reporting.shared import get_image_pixel_size, load_json

REGISTER_HEADING = "REGISTRO SEMANAL DE TRABAJOS EFECTUADOS"
PEA_HEADING = "PLAN ESPECIFICO DE APRENDIZAJE (PEA)"
REPORT_HEADING = "INFORME DE FORMACION PRACTICA EN EMPRESA"
DIAGRAM_TABLE_HEADING = "HACER ESQUEMA, DIBUJO O DIAGRAMA"
EVALUATION_TABLE_HEADING = "EVALUACION DEL INFORME POR EL INSTRUCTOR"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ASCII_OUTPUT_FALLBACK = False

HEADER_FONT = "JetBrains Mono ExtraBold"
HEADER_SIZE = 12
BODY_FONT = "Iosevka NF"
BODY_SIZE = 9
TABLE_FONT = "JetBrains Mono ExtraBold"
TABLE_SIZE = 9
PEA_TEXT_FONT = "Iosevka NF"
PEA_META_FONT = "Calibri"

SMART_CHAR_MAP = str.maketrans(
    {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u2026": "...",
        "\u00a0": " ",
        "\u200b": "",
        "\ufeff": "",
    }
)


def normalize_context(data: dict) -> dict:
    context = dict(data)
    context.setdefault("rotaciones", [])
    context.setdefault("pea_items", [])
    context.setdefault("semanas", [])
    add_week_blocks(context)
    for key, value in list(context.items()):
        if value is None:
            context[key] = ""
    return context


def normalize_output_string(text: str) -> str:
    value = unicodedata.normalize("NFC", str(text)).translate(SMART_CHAR_MAP)
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    if not ASCII_OUTPUT_FALLBACK:
        return value
    simplified = unicodedata.normalize("NFKD", value)
    simplified = "".join(ch for ch in simplified if not unicodedata.combining(ch))
    return simplified.encode("ascii", "ignore").decode("ascii")


def normalize_lookup_text(value: str) -> str:
    simplified = unicodedata.normalize("NFKD", str(value or ""))
    simplified = "".join(ch for ch in simplified if not unicodedata.combining(ch))
    return " ".join(simplified.upper().split())


def text_matches(value: str, target: str) -> bool:
    return normalize_lookup_text(value) == normalize_lookup_text(target)


def text_startswith(value: str, prefix: str) -> bool:
    return normalize_lookup_text(value).startswith(normalize_lookup_text(prefix))


def text_contains(value: str, target: str) -> bool:
    return normalize_lookup_text(target) in normalize_lookup_text(value)


def normalize_output_data(value, key: str | None = None):
    if isinstance(value, dict):
        result = {}
        for child_key, child_value in value.items():
            if child_key.endswith("_path"):
                result[child_key] = child_value
            else:
                result[child_key] = normalize_output_data(child_value, child_key)
        return result
    if isinstance(value, list):
        return [normalize_output_data(item, key) for item in value]
    if isinstance(value, str):
        return normalize_output_string(value)
    return value


def blank_day() -> dict:
    return {"fecha": "", "actividades": "", "horas": ""}


def blank_week() -> dict:
    return {
        "numero_semana": "",
        "lunes": blank_day(),
        "martes": blank_day(),
        "miercoles": blank_day(),
        "jueves": blank_day(),
        "viernes": blank_day(),
        "sabado": blank_day(),
        "horas_totales": "",
    }


def normalize_day(day: dict | None) -> dict:
    base = blank_day()
    if not day:
        return base
    for key in base:
        value = day.get(key, "")
        base[key] = "" if value is None else value
    return base


def build_week_lookup(week: dict | None) -> dict:
    result = blank_week()
    if not week:
        return result
    result["numero_semana"] = week.get("numero_semana", "") or ""
    result["horas_totales"] = week.get("horas_totales", "") or ""
    by_name = {}
    for day in week.get("dias", []):
        by_name[normalize_lookup_text(day.get("dia", ""))] = day
    for source, target in {
        "LUNES": "lunes",
        "MARTES": "martes",
        "MIERCOLES": "miercoles",
        "JUEVES": "jueves",
        "VIERNES": "viernes",
        "SABADO": "sabado",
    }.items():
        if source in by_name:
            result[target] = normalize_day(by_name[source])
    return result


def add_week_blocks(context: dict) -> None:
    normalized = [build_week_lookup(week) for week in context.get("semanas", [])[:3]]
    while len(normalized) < 3:
        normalized.append(blank_week())
    context["semana_1"], context["semana_2"], context["semana_3"] = normalized
    context.setdefault("nombre_empresa", "")


def render_docx(template_path: Path, data_path: Path, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_BREAK
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docxtpl import DocxTemplate, InlineImage, Listing
    except ImportError as exc:
        raise SystemExit("Faltan dependencias. Instala el entorno con: uv sync") from exc

    context = normalize_output_data(normalize_context(load_json(data_path, {})))
    postprocess_context = copy.deepcopy(context)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxTemplate(str(template_path))
    add_inline_images(doc, context, InlineImage, Mm)
    add_multiline_fields(context, Listing)
    doc.render(context, autoescape=False)
    doc.save(str(output_path))
    cleanup_rendered_docx(
        output_path,
        Document,
        Paragraph,
        Table,
        WD_BREAK,
        WD_ALIGN_PARAGRAPH,
        WD_CELL_VERTICAL_ALIGNMENT,
        postprocess_context,
    )


def bounded_image_size_mm(path: Path, max_width_mm: float, max_height_mm: float) -> tuple[float, float] | None:
    size = get_image_pixel_size(path)
    if not size:
        return None

    width_px, height_px = size
    if width_px <= 0 or height_px <= 0:
        return None

    scale = min(max_width_mm / width_px, max_height_mm / height_px, 1.0)
    return width_px * scale, height_px * scale


def add_inline_images(doc, context: dict, inline_image_cls, mm_cls) -> None:
    for target_key, path_key, max_width_mm, max_height_mm in (
        ("imagen_esquema_diagrama", "diagrama_path", 105, 70),
        ("firma_estudiante", "firma_estudiante_path", 35, 18),
        ("firma_monitor", "firma_monitor_path", 35, 18),
    ):
        image_path = context.get(path_key)
        if image_path:
            path = Path(image_path)
            if path.exists():
                bounded_size = bounded_image_size_mm(path, max_width_mm, max_height_mm)
                if bounded_size:
                    width_mm, height_mm = bounded_size
                    context[target_key] = inline_image_cls(
                        doc,
                        str(path),
                        width=mm_cls(width_mm),
                        height=mm_cls(height_mm),
                    )
                else:
                    context[target_key] = inline_image_cls(doc, str(path), width=mm_cls(max_width_mm))
                continue
        context.setdefault(target_key, "")


def make_listing(value, listing_cls, bullet: bool = False):
    if value is None:
        return ""
    if isinstance(value, list):
        lines = [str(item).strip() for item in value if str(item).strip()]
        text = "\n".join(f"- {line}" for line in lines) if bullet else "\n".join(lines)
    else:
        text = str(value)
    return listing_cls(text) if text.strip() else ""


def add_multiline_fields(context: dict, listing_cls) -> None:
    for key, bullet in {
        "porque_tarea_significativa": False,
        "descripcion_tarea_significativa": False,
        "maquinas_usadas": True,
        "equipos_usados": True,
        "herramientas_usadas": True,
        "materiales_usados": True,
        "charlas_seguridad": True,
        "explicacion_actividades_y_medidas_seguras_usadas": True,
        "resultado_ejecucion": False,
        "objetivo_logrado_o_no": False,
        "recomendaciones_resultado": True,
    }.items():
        if key in context:
            context[key] = make_listing(context.get(key), listing_cls, bullet=bullet)
    for week_key in ("semana_1", "semana_2", "semana_3"):
        week = context.get(week_key, {})
        for day_key in ("lunes", "martes", "miercoles", "jueves", "viernes", "sabado"):
            day = week.get(day_key, {})
            day["actividades"] = make_listing(day.get("actividades", ""), listing_cls)


def cleanup_rendered_docx(
    output_path: Path,
    document_cls,
    paragraph_cls,
    table_cls,
    break_enum,
    paragraph_align_enum,
    vertical_align_enum,
    source_context: dict,
) -> None:
    doc = document_cls(str(output_path))
    remove_trailing_empty_week_sections(doc, paragraph_cls, table_cls, source_context)
    remove_page_break_from_heading(doc, PEA_HEADING)
    ensure_page_break_before_heading(doc, REGISTER_HEADING, paragraph_cls, table_cls, break_enum, occurrence="first")
    ensure_page_break_before_heading(doc, REPORT_HEADING, paragraph_cls, table_cls, break_enum, occurrence="last")
    ensure_page_break_before_table(doc, DIAGRAM_TABLE_HEADING, paragraph_cls, table_cls, break_enum)
    ensure_page_break_before_table(doc, EVALUATION_TABLE_HEADING, paragraph_cls, table_cls, break_enum)
    repair_dynamic_text(doc, source_context)
    center_diagram_table_content(doc, paragraph_align_enum, vertical_align_enum)
    doc.save(str(output_path))


def iter_block_items(doc, paragraph_cls, table_cls):
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield paragraph_cls(child, doc)
        elif tag == "tbl":
            yield table_cls(child, doc)


def text_value(value, bullet: bool = False) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        items = [str(item).strip() for item in value if str(item).strip()]
        if bullet:
            items = [item if item.startswith("- ") else f"- {item}" for item in items]
        return "\n".join(items)
    return str(value)


def style_runs(runs, font_name: str, size: int, *, bold: bool | None = False) -> None:
    for run in runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold


def style_paragraph(
    paragraph,
    font_name: str,
    size: int,
    *,
    bold: bool | None = False,
    space_before: int = 0,
    space_after: int = 0,
    line_spacing: float = 1.0,
) -> None:
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    runs = paragraph.runs
    if not runs and paragraph.text:
        runs = [paragraph.add_run("")]
    style_runs(runs, font_name, size, bold=bold)


def write_paragraph_text(
    paragraph,
    value: str,
    font_name: str,
    size: int,
    *,
    bold: bool | None = False,
    space_before: int = 0,
    space_after: int = 0,
    line_spacing: float = 1.0,
) -> None:
    paragraph.text = value
    style_paragraph(
        paragraph,
        font_name,
        size,
        bold=bold,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
    )


def write_cell_text(cell, value: str, font_name: str, size: int, *, bold: bool | None = False) -> None:
    cell.text = value
    for paragraph in cell.paragraphs:
        style_paragraph(paragraph, font_name, size, bold=bold, space_before=0, space_after=0, line_spacing=1.0)


def remove_block(block) -> None:
    element = block._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def has_week_content(week: dict | None) -> bool:
    if not week:
        return False
    if str(week.get("numero_semana", "")).strip() or str(week.get("horas_totales", "")).strip():
        return True
    for day_key in ("lunes", "martes", "miercoles", "jueves", "viernes", "sabado"):
        day = week.get(day_key, {}) or {}
        if any(str(day.get(field, "")).strip() for field in ("fecha", "actividades", "horas")):
            return True
    return False


def remove_trailing_empty_week_sections(doc, paragraph_cls, table_cls, source_context: dict) -> None:
    expected_weeks = sum(
        1 for key in ("semana_1", "semana_2", "semana_3") if has_week_content(source_context.get(key))
    )
    blocks = list(iter_block_items(doc, paragraph_cls, table_cls))
    week_start_indices = [
        index
        for index, block in enumerate(blocks)
        if isinstance(block, paragraph_cls) and text_startswith(block.text, "SEMANA")
    ]

    if len(week_start_indices) <= expected_weeks:
        return

    start = week_start_indices[expected_weeks]
    report_index = None
    for index in range(start + 1, len(blocks)):
        block = blocks[index]
        if isinstance(block, paragraph_cls) and text_matches(block.text, REPORT_HEADING):
            report_index = index
            break

    if report_index is None or start >= report_index:
        return

    while start > 0:
        previous = blocks[start - 1]
        if not isinstance(previous, paragraph_cls) or normalize_lookup_text(previous.text):
            break
        start -= 1

    for block in reversed(blocks[start:report_index]):
        remove_block(block)


def paragraph_has_page_break(paragraph) -> bool:
    xml = paragraph._element.xml
    return bool(paragraph.paragraph_format.page_break_before) or 'w:type="page"' in xml


def remove_page_break_from_heading(doc, heading_text: str) -> None:
    for paragraph in doc.paragraphs:
        if text_matches(paragraph.text, heading_text):
            remove_page_break_from_paragraph(paragraph)
            return


def remove_page_break_from_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        for child in list(run._element):
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "br" and child.get(f"{{{W_NS}}}type") == "page":
                run._element.remove(child)
            elif tag == "lastRenderedPageBreak":
                run._element.remove(child)


def ensure_page_break_before_heading(
    doc,
    heading_text: str,
    paragraph_cls,
    table_cls,
    break_enum,
    occurrence: str = "first",
) -> None:
    blocks = list(iter_block_items(doc, paragraph_cls, table_cls))
    candidates = []

    for index, block in enumerate(blocks):
        if isinstance(block, paragraph_cls) and text_matches(block.text, heading_text):
            candidates.append((index, block))

    if not candidates:
        return

    index, block = candidates[-1] if occurrence == "last" else candidates[0]

    if paragraph_has_page_break(block):
        return

    cursor = index - 1
    while cursor >= 0:
        previous = blocks[cursor]
        if not isinstance(previous, paragraph_cls):
            break
        if paragraph_has_page_break(previous):
            return
        if normalize_lookup_text(previous.text):
            break
        cursor -= 1

    page_break_paragraph = block.insert_paragraph_before()
    page_break_paragraph.add_run().add_break(break_enum.PAGE)


def ensure_page_break_before_table(doc, table_text: str, paragraph_cls, table_cls, break_enum) -> None:
    blocks = list(iter_block_items(doc, paragraph_cls, table_cls))
    target_index = None

    for index, block in enumerate(blocks):
        if isinstance(block, table_cls) and text_contains(block._element.xpath("string(.)"), table_text):
            target_index = index
            break

    if target_index is None:
        return

    cursor = target_index - 1
    anchor = None
    blank_paragraphs = []
    while cursor >= 0:
        previous = blocks[cursor]
        if not isinstance(previous, paragraph_cls):
            break
        if paragraph_has_page_break(previous):
            anchor = previous
            break
        if normalize_lookup_text(previous.text):
            anchor = previous
            break
        blank_paragraphs.append(previous)
        anchor = previous
        cursor -= 1

    if anchor is None:
        return

    for paragraph in blank_paragraphs[1:]:
        remove_block(paragraph)

    if not paragraph_has_page_break(anchor):
        anchor.add_run().add_break(break_enum.PAGE)


def repair_dynamic_text(doc, context: dict) -> None:
    repair_header_paragraphs(doc, context)
    repair_report_paragraphs(doc, context)
    repair_rotation_table(doc, context)
    repair_pea_table(doc, context)
    repair_week_tables(doc, context)


def center_diagram_table_content(doc, paragraph_align_enum, vertical_align_enum) -> None:
    table = first_table_matching(doc, DIAGRAM_TABLE_HEADING)
    if table is None:
        return

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            continue
        for cell in row.cells:
            cell.vertical_alignment = vertical_align_enum.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = paragraph_align_enum.CENTER


def repair_header_paragraphs(doc, context: dict) -> None:
    replacements = {
        "CFP/UCP/ESCUELA:": f"CFP/UCP/ESCUELA: __{context.get('escuela', '')}___",
        "ESTUDIANTE:": f"ESTUDIANTE: ___{context.get('nombre_estudiante', '')}_______________________",
        "ID:": (
            f"ID: ___{context.get('id_estudiante', '')}________ BLOQUE: "
            f"{context.get('bloque', '')}____________________"
        ),
        "CARRERA:": f"CARRERA: {context.get('carrera', '')}____",
        "INSTRUCTOR:": f"INSTRUCTOR: ___{context.get('instructor', '')}_________________________",
        "SEMESTRE:": (
            f"SEMESTRE: ___{context.get('semestre', '')}___ DEL: "
            f"___{context.get('fecha_inicio_semestre', '')}___ AL: "
            f"___{context.get('fecha_fin_semestre', '')}_______________"
        ),
    }

    for paragraph in doc.paragraphs:
        for prefix, value in replacements.items():
            if text_startswith(paragraph.text, prefix):
                write_paragraph_text(paragraph, value, HEADER_FONT, HEADER_SIZE, bold=False)
                break


def find_paragraph_index(doc, prefix: str) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        if text_startswith(paragraph.text, prefix):
            return index
    return None


def section_slots(paragraphs, start_index: int, stop_prefixes: list[str], count: int) -> list:
    filled = []
    blank = []
    for paragraph in paragraphs[start_index + 1 :]:
        if any(text_startswith(paragraph.text, prefix) for prefix in stop_prefixes):
            break
        if normalize_lookup_text(paragraph.text):
            filled.append(paragraph)
        else:
            blank.append(paragraph)
    return (filled + blank)[:count]


def repair_report_paragraphs(doc, context: dict) -> None:
    for paragraph in doc.paragraphs:
        if text_startswith(paragraph.text, "TAREA:"):
            write_paragraph_text(paragraph, f"Tarea: {text_value(context.get('tarea_significativa'))}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "JUSTIFICACION:"):
            write_paragraph_text(paragraph, f"Justificacion: {text_value(context.get('porque_tarea_significativa'))}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "MAQUINAS:"):
            write_paragraph_text(paragraph, f"Maquinas: {text_value(context.get('maquinas_usadas'), bullet=True)}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "EQUIPOS:"):
            write_paragraph_text(paragraph, f"Equipos: {text_value(context.get('equipos_usados'), bullet=True)}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "HERRAMIENTAS:"):
            write_paragraph_text(paragraph, f"Herramientas: {text_value(context.get('herramientas_usadas'), bullet=True)}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "MATERIALES:"):
            write_paragraph_text(paragraph, f"Materiales: {text_value(context.get('materiales_usados'), bullet=True)}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "RESULTADO:"):
            write_paragraph_text(paragraph, f"Resultado: {text_value(context.get('resultado_ejecucion'))}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "OBJETIVO:"):
            write_paragraph_text(paragraph, f"Objetivo: {text_value(context.get('objetivo_logrado_o_no'))}", BODY_FONT, BODY_SIZE)
        elif text_startswith(paragraph.text, "RECOMENDACIONES:"):
            write_paragraph_text(
                paragraph,
                f"Recomendaciones: {text_value(context.get('recomendaciones_resultado'), bullet=True)}",
                BODY_FONT,
                BODY_SIZE,
            )

    paragraphs = doc.paragraphs
    process_index = find_paragraph_index(doc, "DESCRIPCIÓN DEL PROCESO:")
    if process_index is not None:
        slots = section_slots(paragraphs, process_index, ["MÁQUINAS, EQUIPOS, HERRAMIENTAS Y MATERIALES"], 1)
        if slots:
            write_paragraph_text(slots[0], text_value(context.get("descripcion_tarea_significativa")), BODY_FONT, BODY_SIZE)

    safety_index = find_paragraph_index(doc, "SEGURIDAD E HIGIENE INDUSTRIAL/AMBIENTAL")
    if safety_index is not None:
        slots = section_slots(paragraphs, safety_index, ["RESULTADOS DE LA EJECUCIÓN DE LA TAREA/RECOMENDACIONES"], 2)
        if slots:
            write_paragraph_text(slots[0], text_value(context.get("charlas_seguridad"), bullet=True), BODY_FONT, BODY_SIZE)
        if len(slots) > 1:
            write_paragraph_text(
                slots[1],
                text_value(context.get("explicacion_actividades_y_medidas_seguras_usadas"), bullet=True),
                BODY_FONT,
                BODY_SIZE,
            )


def first_table_matching(doc, needle: str):
    for table in doc.tables:
        if text_contains(table._element.xpath("string(.)"), needle):
            return table
    return None


def repair_rotation_table(doc, context: dict) -> None:
    table = first_table_matching(doc, "CUADRO DE ROTACIONES")
    if table is None:
        return

    for row_index, item in enumerate(context.get("rotaciones", []), start=3):
        if row_index >= len(table.rows):
            break
        row = table.rows[row_index]
        write_cell_text(row.cells[0], text_value(item.get("area")), TABLE_FONT, TABLE_SIZE, bold=False)
        write_cell_text(row.cells[1], text_value(item.get("desde")), TABLE_FONT, TABLE_SIZE, bold=False)
        write_cell_text(row.cells[2], text_value(item.get("hasta")), TABLE_FONT, TABLE_SIZE, bold=False)
        write_cell_text(row.cells[3], text_value(item.get("semana")), TABLE_FONT, TABLE_SIZE, bold=False)


def repair_pea_table(doc, context: dict) -> None:
    table = first_table_matching(doc, "OPERACIONES/TAREAS")
    if table is None:
        return

    for row_index, item in enumerate(context.get("pea_items", []), start=2):
        if row_index >= len(table.rows):
            break
        row = table.rows[row_index]
        values = [
            text_value(item.get("numero")),
            text_value(item.get("descripcion")),
            text_value(item.get("op1")),
            text_value(item.get("op2")),
            text_value(item.get("op3")),
            text_value(item.get("op4")),
            "",
        ]
        write_cell_text(row.cells[0], values[0], PEA_META_FONT, BODY_SIZE, bold=False)
        write_cell_text(row.cells[1], values[1], PEA_TEXT_FONT, BODY_SIZE, bold=False)
        for cell, value in zip(row.cells[2:], values[2:]):
            write_cell_text(cell, value, PEA_META_FONT, BODY_SIZE, bold=False)


def repair_week_tables(doc, context: dict) -> None:
    tables = [
        table
        for table in doc.tables
        if table.rows and table.rows[0].cells and text_matches(table.rows[0].cells[0].text, "DIA")
    ]
    weeks = [
        week
        for week in (context.get("semana_1"), context.get("semana_2"), context.get("semana_3"))
        if has_week_content(week)
    ]
    day_keys = [
        ("lunes", "LUNES"),
        ("martes", "MARTES"),
        ("miercoles", "MIÉRCOLES"),
        ("jueves", "JUEVES"),
        ("viernes", "VIERNES"),
        ("sabado", "SÁBADO"),
    ]

    for table, week in zip(tables, weeks):
        for row_index, (day_key, day_label) in enumerate(day_keys, start=1):
            if row_index >= len(table.rows):
                break
            day = week.get(day_key, {})
            row = table.rows[row_index]
            write_cell_text(row.cells[0], f"{day_label}\n{day.get('fecha', '')}".strip(), TABLE_FONT, TABLE_SIZE, bold=False)
            write_cell_text(row.cells[1], text_value(day.get("actividades")), TABLE_FONT, TABLE_SIZE, bold=False)
            write_cell_text(row.cells[2], text_value(day.get("horas")), TABLE_FONT, TABLE_SIZE, bold=False)

        if len(table.rows) > 7:
            write_cell_text(table.rows[7].cells[1], "TOTAL", TABLE_FONT, TABLE_SIZE, bold=False)
            write_cell_text(table.rows[7].cells[2], text_value(week.get("horas_totales")), TABLE_FONT, TABLE_SIZE, bold=False)


def main() -> None:
    parser = argparse.ArgumentParser(description="Renderiza un informe FPE desde una plantilla DOCX y un JSON.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--data", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    render_docx(args.template, args.data, args.output)


if __name__ == "__main__":
    main()
